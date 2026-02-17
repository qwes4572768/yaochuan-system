"""級距表：列表、依計算月份有效版本、匯入 JSON/Excel"""
import json
from datetime import date
from decimal import Decimal
from io import BytesIO
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app import crud
from app.crud import RATE_TABLE_TYPES, get_effective_rate_table
from app.models import RateTable, RateItem
from app.schemas import RateTableRead, RateItemRead, RateTableImportPayload, RateTableImportTable, RateTableImportItem

router = APIRouter(prefix="/api/rate-tables", tags=["rate-tables"])


def _dict_to_import_table(d: dict) -> RateTableImportTable:
    """JSON 解析後轉成 RateTableImportTable"""
    items = d.get("items", [])
    return RateTableImportTable(
        type=d["type"],
        version=d.get("version", "2025-01"),
        effective_from=d.get("effective_from", "2025-01-01"),
        effective_to=d.get("effective_to"),
        total_rate=d.get("total_rate"),
        note=d.get("note"),
        items=[RateTableImportItem(**it) if isinstance(it, dict) else it for it in items],
    )


def _rate_table_to_read(t: RateTable) -> RateTableRead:
    return RateTableRead(
        id=t.id,
        type=t.type,
        version=t.version,
        effective_from=t.effective_from,
        effective_to=t.effective_to,
        total_rate=t.total_rate,
        note=t.note,
        items=[
            RateItemRead(
                id=it.id,
                table_id=it.table_id,
                level_name=it.level_name,
                salary_min=it.salary_min,
                salary_max=it.salary_max,
                insured_salary=it.insured_salary,
                employee_rate=it.employee_rate,
                employer_rate=it.employer_rate,
                gov_rate=it.gov_rate,
                fixed_amount_if_any=it.fixed_amount_if_any,
            )
            for it in t.items
        ],
    )


@router.get("", response_model=List[RateTableRead])
async def list_rate_tables(
    type: Optional[str] = Query(None, description="篩選類型"),
    db: AsyncSession = Depends(get_db),
):
    """列出級距表（可篩選 type）"""
    if type and type not in RATE_TABLE_TYPES:
        raise HTTPException(status_code=400, detail=f"type 須為: {list(RATE_TABLE_TYPES)}")
    tables = await crud.list_rate_tables(db, table_type=type)
    return [_rate_table_to_read(t) for t in tables]


@router.get("/effective", response_model=Dict[str, Any])
async def get_effective_for_month(
    year: int = Query(..., description="計算年度"),
    month: int = Query(..., ge=1, le=12, description="計算月份"),
    db: AsyncSession = Depends(get_db),
):
    """依計算月份取得當月有效之各類型級距表（供試算套用）"""
    as_of = date(year, month, 1)
    result = {}
    for t in RATE_TABLE_TYPES:
        tbl = await get_effective_rate_table(db, t, as_of)
        if tbl:
            result[t] = _rate_table_to_read(tbl)
        else:
            result[t] = None
    return result


@router.post("/import", response_model=List[RateTableRead])
async def import_rate_tables(
    body: Optional[RateTableImportPayload] = None,
    file: Optional[UploadFile] = None,
    db: AsyncSession = Depends(get_db),
):
    """
    匯入級距表。二擇一：
    1) body: JSON { "tables": [ { type, version, effective_from, effective_to?, total_rate?, note?, items: [...] }, ... ] }
    2) file: 上傳 JSON 或 Excel（.xlsx）檔，格式同 body.tables
    """
    if body is None and file is None:
        raise HTTPException(status_code=400, detail="請提供 body 或 file")
    if body is not None and file is not None:
        raise HTTPException(status_code=400, detail="請擇一提供 body 或 file")

    tables_data: List[RateTableImportTable]
    if body is not None:
        tables_data = body.tables
    else:
        raw = await file.read()
        fn = (file.filename or "").lower()
        try:
            if fn.endswith(".json"):
                try:
                    data = json.loads(raw.decode("utf-8"))
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"JSON 解析失敗: {e}")
                tables_data = data.get("tables", data) if isinstance(data, dict) else data
                if not isinstance(tables_data, list):
                    tables_data = [tables_data]
                tables_data = [_dict_to_import_table(x) for x in tables_data]
            elif fn.endswith(".xlsx"):
                tables_data = _parse_excel_rate_tables(BytesIO(raw))
            elif fn.endswith(".docx"):
                tables_data = _parse_docx_rate_tables(BytesIO(raw))
            else:
                raise HTTPException(status_code=400, detail="僅支援 .json、.xlsx 或 .docx")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"檔案解析失敗: {e}")

    created_ids: List[int] = []
    for pt in tables_data:
        if pt.type not in RATE_TABLE_TYPES:
            raise HTTPException(status_code=400, detail=f"不支援的 type: {pt.type}")
        eff_from = date.fromisoformat(str(pt.effective_from)[:10])
        eff_to = date.fromisoformat(str(pt.effective_to)[:10]) if pt.effective_to else None
        total_rate = Decimal(str(pt.total_rate)) if pt.total_rate is not None else None
        tbl = RateTable(
            type=pt.type,
            version=pt.version,
            effective_from=eff_from,
            effective_to=eff_to,
            total_rate=total_rate,
            note=pt.note,
        )
        db.add(tbl)
        await db.flush()
        for it in pt.items:
            salary_min = Decimal(str(it.salary_min))
            salary_max = Decimal(str(it.salary_max))
            insured = Decimal(str(it.insured_salary)) if it.insured_salary is not None else None
            emp_r = Decimal(str(it.employee_rate))
            emp_r_company = Decimal(str(it.employer_rate))
            gov_r = Decimal(str(it.gov_rate)) if it.gov_rate is not None else None
            fixed = Decimal(str(it.fixed_amount_if_any)) if it.fixed_amount_if_any is not None else None
            db.add(
                RateItem(
                    table_id=tbl.id,
                    level_name=it.level_name,
                    salary_min=salary_min,
                    salary_max=salary_max,
                    insured_salary=insured,
                    employee_rate=emp_r,
                    employer_rate=emp_r_company,
                    gov_rate=gov_r,
                    fixed_amount_if_any=fixed,
                )
            )
        await db.flush()
        created_ids.append(tbl.id)
    await db.commit()
    # 重新載入含 items 的完整資料
    result = []
    for tid in created_ids:
        tbl = await crud.get_rate_table_by_id(db, tid)
        if tbl:
            result.append(_rate_table_to_read(tbl))
    return result


def _parse_excel_rate_tables(stream: BytesIO) -> List[RateTableImportTable]:
    """簡易解析 Excel：每 sheet 為一類型，欄位 type, version, effective_from, effective_to, total_rate, note；其後為 items 表頭 salary_min, salary_max, insured_salary, employee_rate, employer_rate, gov_rate"""
    from openpyxl import load_workbook
    wb = load_workbook(stream, read_only=True, data_only=True)
    tables_data: List[RateTableImportTable] = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c).strip().lower() if c else "" for c in rows[0]]
        header_raw = [str(c).strip() if c else "" for c in rows[0]]
        # 找 type 或 類型 欄
        type_idx = next((i for i, h in enumerate(header) if h == "type"), None)
        if type_idx is None:
            type_idx = next((i for i, r in enumerate(header_raw) if "類型" in r or r == "類型"), None)
        if type_idx is None:
            continue
        data_rows = rows[1:]
        # 第一列為表頭資訊（meta）
        meta = {}
        for i, h in enumerate(header):
            if i < len(data_rows) and data_rows[0][i] is not None:
                meta[h] = data_rows[0][i]
        for i, r in enumerate(header_raw):
            if i < len(data_rows) and data_rows[0][i] is not None and (r and r not in meta):
                meta[r] = data_rows[0][i]
        type_val = str(meta.get("type") or meta.get("類型") or "").strip()
        if type_val not in RATE_TABLE_TYPES:
            continue
        version = str(meta.get("version", "2025-01"))
        eff_from = str(meta.get("effective_from", "2025-01-01"))[:10]
        eff_to = meta.get("effective_to")
        if eff_to is not None:
            eff_to = str(eff_to)[:10]
        total_rate = meta.get("total_rate")
        note = meta.get("note")
        # items: 從第 2 列起，欄位 salary_min, salary_max, insured_salary, employee_rate, employer_rate, gov_rate（支援英文與中文表頭）
        item_header = [str(c).strip().lower() if c else "" for c in (data_rows[1] if len(data_rows) > 1 else [])]
        item_header_raw = [str(c).strip() if c else "" for c in (data_rows[1] if len(data_rows) > 1 else [])]
        salary_min_i = next((i for i, h in enumerate(item_header) if "salary_min" in h or h == "min"), None)
        if salary_min_i is None:
            salary_min_i = next((i for i, r in enumerate(item_header_raw) if "下限" in r), 0)
        salary_max_i = next((i for i, h in enumerate(item_header) if "salary_max" in h or h == "max"), None)
        if salary_max_i is None:
            salary_max_i = next((i for i, r in enumerate(item_header_raw) if "上限" in r), 1)
        insured_i = next((i for i, h in enumerate(item_header) if "insured" in h or "level" in h), None)
        if insured_i is None:
            insured_i = next((i for i, r in enumerate(item_header_raw) if "投保" in r or "級距" in r), 2)
        emp_i = next((i for i, h in enumerate(item_header) if "employee" in h), None)
        if emp_i is None:
            emp_i = next((i for i, r in enumerate(item_header_raw) if "個人" in r), 3)
        comp_i = next((i for i, h in enumerate(item_header) if "employer" in h or "company" in h), None)
        if comp_i is None:
            comp_i = next((i for i, r in enumerate(item_header_raw) if "公司" in r), 4)
        gov_i = next((i for i, h in enumerate(item_header) if "gov" in h), None)
        if gov_i is None:
            gov_i = next((i for i, r in enumerate(item_header_raw) if "政府" in r), 5)
        # 若仍為 None 則用預設欄位索引
        if salary_min_i is None:
            salary_min_i = 0
        if salary_max_i is None:
            salary_max_i = 1
        if insured_i is None:
            insured_i = 2
        if emp_i is None:
            emp_i = 3
        if comp_i is None:
            comp_i = 4
        if gov_i is None:
            gov_i = 5
        items = []
        for row in data_rows[2:]:
            if row is None or all(c is None for c in row):
                continue
            try:
                salary_min = int(float(row[salary_min_i])) if row[salary_min_i] is not None else 0
                salary_max = int(float(row[salary_max_i])) if row[salary_max_i] is not None else 0
                insured = int(float(row[insured_i])) if insured_i < len(row) and row[insured_i] is not None else None
                emp_r = float(row[emp_i]) if emp_i < len(row) and row[emp_i] is not None else 0
                comp_r = float(row[comp_i]) if comp_i < len(row) and row[comp_i] is not None else 0
                gov_r = float(row[gov_i]) if gov_i < len(row) and row[gov_i] is not None else None
                items.append(
                    RateTableImportItem(
                        level_name=str(insured) if insured else None,
                        salary_min=salary_min,
                        salary_max=salary_max,
                        insured_salary=insured,
                        employee_rate=emp_r,
                        employer_rate=comp_r,
                        gov_rate=gov_r,
                    )
                )
            except (ValueError, TypeError, IndexError):
                continue
        if not items and type_val == "labor_insurance":
            # 若無 items 可略過或給一筆預設
            items = [
                RateTableImportItem(
                    level_name="26400",
                    salary_min=0,
                    salary_max=999999,
                    insured_salary=26400,
                    employee_rate=0.2,
                    employer_rate=0.7,
                    gov_rate=0.1,
                )
            ]
        tables_data.append(
            RateTableImportTable(
                type=type_val,
                version=version,
                effective_from=eff_from,
                effective_to=eff_to,
                total_rate=float(total_rate) if total_rate is not None else None,
                note=str(note) if note else None,
                items=items,
            )
        )
    return tables_data


def _parse_docx_rate_tables(stream: BytesIO) -> List[RateTableImportTable]:
    """解析 Word .docx：每個表格視為一類型，第一列為表頭（type/version/effective_from 或級距欄位），其後為級距列。"""
    from docx import Document
    doc = Document(stream)
    tables_data: List[RateTableImportTable] = []
    type_order = list(RATE_TABLE_TYPES)
    for ti, table in enumerate(doc.tables):
        rows = [[(c.text or "").strip() for c in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [h.lower() for h in rows[0]]
        has_type_col = any(h == "type" or "類型" in h for h in header)
        if has_type_col and len(rows) > 1:
            meta = {header[i]: rows[1][i] if i < len(rows[1]) else "" for i in range(len(header))}
            data_start = 2
        else:
            meta = {header[i]: rows[0][i] if i < len(rows[0]) else "" for i in range(len(header))}
            data_start = 1
        type_val = (meta.get("type") or meta.get("類型") or "").strip()
        if not type_val and ti < len(type_order):
            type_val = type_order[ti]
        if type_val not in RATE_TABLE_TYPES:
            continue
        version = str(meta.get("version", "2025-01"))
        eff_from = str(meta.get("effective_from", "2025-01-01"))[:10]
        eff_to = meta.get("effective_to")
        if eff_to is not None:
            eff_to = str(eff_to)[:10]
        total_rate = meta.get("total_rate")
        note = meta.get("note")
        item_header = [h.lower() for h in (rows[data_start] if data_start < len(rows) else [])]
        salary_min_i = next((i for i, h in enumerate(item_header) if "salary_min" in h or h == "min" or "下限" in h), 0)
        salary_max_i = next((i for i, h in enumerate(item_header) if "salary_max" in h or h == "max" or "上限" in h), 1)
        insured_i = next((i for i, h in enumerate(item_header) if "insured" in h or "level" in h or "投保" in h or "級距" in h), 2)
        emp_i = next((i for i, h in enumerate(item_header) if "employee" in h or "個人" in h), 3)
        comp_i = next((i for i, h in enumerate(item_header) if "employer" in h or "company" in h or "公司" in h), 4)
        gov_i = next((i for i, h in enumerate(item_header) if "gov" in h or "政府" in h), 5)
        items: List[RateTableImportItem] = []
        for row in rows[data_start + 1:]:
            if not row or all(not c for c in row):
                continue
            try:
                salary_min = int(float(row[salary_min_i])) if salary_min_i < len(row) and row[salary_min_i] else 0
                salary_max = int(float(row[salary_max_i])) if salary_max_i < len(row) and row[salary_max_i] else 0
                insured = None
                if insured_i < len(row) and row[insured_i]:
                    try:
                        insured = int(float(row[insured_i]))
                    except (ValueError, TypeError):
                        pass
                emp_r = float(row[emp_i]) if emp_i < len(row) and row[emp_i] else 0
                comp_r = float(row[comp_i]) if comp_i < len(row) and row[comp_i] else 0
                gov_r = float(row[gov_i]) if gov_i < len(row) and row[gov_i] else None
                items.append(
                    RateTableImportItem(
                        level_name=str(insured) if insured else None,
                        salary_min=salary_min,
                        salary_max=salary_max,
                        insured_salary=insured,
                        employee_rate=emp_r,
                        employer_rate=comp_r,
                        gov_rate=gov_r,
                    )
                )
            except (ValueError, TypeError, IndexError):
                continue
        if not items and type_val == "labor_insurance":
            items = [
                RateTableImportItem(
                    level_name="26400",
                    salary_min=0,
                    salary_max=999999,
                    insured_salary=26400,
                    employee_rate=0.2,
                    employer_rate=0.7,
                    gov_rate=0.1,
                )
            ]
        tables_data.append(
            RateTableImportTable(
                type=type_val,
                version=version,
                effective_from=eff_from,
                effective_to=eff_to,
                total_rate=float(total_rate) if total_rate is not None else None,
                note=str(note) if note else None,
                items=items,
            )
        )
    return tables_data
